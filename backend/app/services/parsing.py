from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from app.schemas.job import ParsedDocument


class ParsingService:
    def parse_files(self, paths: list[Path]) -> list[ParsedDocument]:
        documents: list[ParsedDocument] = []
        for index, path in enumerate(paths, start=1):
            parsed = self.parse_file(path, index)
            if parsed is not None:
                documents.append(parsed)
        return documents

    def parse_file(self, path: Path, index: int) -> ParsedDocument | None:
        suffix = path.suffix.lower()
        try:
            if suffix in {".md", ".txt"}:
                return self._parse_text(path, index)
            if suffix == ".docx":
                return self._parse_docx(path, index)
            if suffix == ".pdf":
                return self._parse_pdf(path, index)
            if suffix == ".xlsx":
                return self._parse_xlsx(path, index)
        except Exception as exc:  # pragma: no cover - best effort parsing
            return ParsedDocument(
                id=f"doc-{index}",
                file_path=str(path),
                file_type=suffix.lstrip('.') or 'unknown',
                title=path.stem,
                excerpt=f"Failed to parse file: {exc}",
                sections=[],
                text_content="",
                metadata={"error": str(exc)},
            )
        return None

    def _parse_text(self, path: Path, index: int) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        sections = [line.strip('# ').strip() for line in text.splitlines() if line.startswith('#')][:12]
        return ParsedDocument(
            id=f"doc-{index}",
            file_path=str(path),
            file_type=path.suffix.lstrip('.'),
            title=sections[0] if sections else path.stem,
            excerpt=text[:240],
            sections=sections,
            text_content=text,
            metadata={"line_count": len(text.splitlines())},
        )

    def _parse_docx(self, path: Path, index: int) -> ParsedDocument:
        document = DocxDocument(path)
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(
            id=f"doc-{index}",
            file_path=str(path),
            file_type="docx",
            title=paragraphs[0][:120] if paragraphs else path.stem,
            excerpt=text[:240],
            sections=paragraphs[:10],
            text_content=text,
            metadata={"paragraph_count": len(paragraphs)},
        )

    def _parse_pdf(self, path: Path, index: int) -> ParsedDocument:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages[:20]:
            pages.append((page.extract_text() or "").strip())
        text = "\n\n".join(filter(None, pages))
        return ParsedDocument(
            id=f"doc-{index}",
            file_path=str(path),
            file_type="pdf",
            title=path.stem,
            excerpt=text[:240],
            sections=[f"page-{i + 1}" for i, page_text in enumerate(pages) if page_text][:10],
            text_content=text,
            metadata={"page_count": len(reader.pages)},
        )

    def _parse_xlsx(self, path: Path, index: int) -> ParsedDocument:
        workbook = load_workbook(filename=path, read_only=True, data_only=True)
        sheet_summaries: list[str] = []
        text_blocks: list[str] = []
        for sheet in workbook.worksheets[:5]:
            rows = []
            for row in sheet.iter_rows(min_row=1, max_row=12, values_only=True):
                values = [str(cell).strip() for cell in row if cell not in (None, "")]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                sheet_summaries.append(sheet.title)
                text_blocks.append(f"[{sheet.title}]\n" + "\n".join(rows))
        text = "\n\n".join(text_blocks)
        return ParsedDocument(
            id=f"doc-{index}",
            file_path=str(path),
            file_type="xlsx",
            title=path.stem,
            excerpt=text[:240],
            sections=sheet_summaries,
            text_content=text,
            metadata={"sheet_count": len(workbook.sheetnames)},
        )
